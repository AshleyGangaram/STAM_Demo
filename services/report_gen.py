"""
PDF and DOCX report generator for STAM analysis reports.
Uses reportlab for PDF and python-docx for Word documents.
"""

from __future__ import annotations

import io
import os
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from models.schemas import AnalysisReport

_BLUE = RGBColor(0x1F, 0x4E, 0x79)    # Gauteng navy
_GREEN = RGBColor(0x1A, 0x74, 0x31)   # Priority green
_GREY = RGBColor(0x60, 0x60, 0x60)    # Body text grey

# ReportLab colors
_RL_BLUE = colors.HexColor("#1F4E79")
_RL_GREEN = colors.HexColor("#1A7431")
_RL_GREY = colors.HexColor("#606060")


def generate_report_pdf(report: AnalysisReport, projects: list) -> bytes:
    """
    Generate a STAM analysis report as a PDF document.
    Returns the document as bytes for Streamlit download.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=1*inch, rightMargin=1*inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=_RL_BLUE,
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=11,
        textColor=_RL_GREY,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=_RL_BLUE,
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        alignment=TA_LEFT
    )

    story = []

    # Cover
    story.append(Paragraph("STAM", title_style))
    story.append(Paragraph("Spatial Transformation Appraisal Mechanism", subtitle_style))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(report.title, heading_style))
    story.append(Paragraph(
        f"Generated: {report.generated_at} | Municipality: {report.municipality} | Sector: {report.sector.title()}",
        subtitle_style
    ))
    story.append(Spacer(1, 0.3*inch))

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(report.executive_summary, body_style))
    story.append(Spacer(1, 0.2*inch))

    # Recommendation
    if report.recommendation:
        rec_style = ParagraphStyle(
            'Recommendation',
            parent=styles['Normal'],
            fontSize=10,
            textColor=_RL_GREEN,
            leftIndent=0.2*inch,
            spaceAfter=8,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(f"STAM Recommendation: {report.recommendation}", rec_style))
        story.append(Spacer(1, 0.2*inch))

    # Sections
    for section in report.sections:
        story.append(Paragraph(section.heading, heading_style))
        story.append(Paragraph(section.content, body_style))
        story.append(Spacer(1, 0.1*inch))

    # Scorecard Table
    if projects:
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("STAM Scorecard Summary", heading_style))

        table_data = [
            ["Project ID", "Project Name", "Type", "Score", "Classification", "Budget (ZAR)"]
        ]
        for p in projects:
            budget = getattr(p, "budget_rands", 0) or 0
            table_data.append([
                str(getattr(p, "project_id", "")),
                str(getattr(p, "name", ""))[:20],
                str(getattr(p, "project_type", "")),
                str(getattr(p, "total_score", 0)),
                str(getattr(p, "classification", "")),
                f"R{budget:,.0f}"
            ])

        table = Table(table_data, colWidths=[0.8*inch, 1.5*inch, 0.7*inch, 0.6*inch, 1*inch, 1.2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), _RL_BLUE),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        story.append(table)

    # Risk notes
    if report.risk_notes:
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Risk and Readiness Notes", heading_style))
        story.append(Paragraph(report.risk_notes, body_style))

    # Footer
    story.append(Spacer(1, 0.3*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=7,
        textColor=_RL_GREY,
        alignment=TA_CENTER
    )
    story.append(Paragraph(
        "Produced by STAM — Spatial Transformation Appraisal Mechanism | TERRA VITAL / Vastpoint | Gauteng Province",
        footer_style
    ))

    doc.build(story)
    return buf.getvalue()


def _set_heading_style(para, level: int = 1, colour: RGBColor = _BLUE):
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.bold = True
    run.font.color.rgb = colour
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def generate_report_docx(report: AnalysisReport, projects: list) -> bytes:
    """
    Generate a STAM analysis report as a Word document.
    Returns the document as bytes for Streamlit download.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover ────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("STAM")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = _BLUE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Spatial Transformation Appraisal Mechanism")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = _GREY
    sub_run.italic = True

    doc.add_paragraph()

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = title2.add_run(report.title)
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = _BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Generated: {report.generated_at}  |  "
        f"Municipality: {report.municipality}  |  "
        f"Sector: {report.sector.title()}"
    ).font.color.rgb = _GREY

    doc.add_paragraph()
    doc.add_paragraph("─" * 70)

    # ── Executive Summary ─────────────────────────────────────────────────────
    h = doc.add_paragraph("Executive Summary")
    _set_heading_style(h, level=1)
    doc.add_paragraph(report.executive_summary)

    # ── Recommendation callout ────────────────────────────────────────────────
    if report.recommendation:
        box = doc.add_paragraph()
        box_run = box.add_run(f"STAM Recommendation:  {report.recommendation}")
        box_run.bold = True
        box_run.font.color.rgb = _GREEN
        box_run.font.size = Pt(12)
        box.paragraph_format.left_indent = Inches(0.3)
        box.paragraph_format.space_before = Pt(6)
        box.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ── Sections ──────────────────────────────────────────────────────────────
    for section in report.sections:
        h = doc.add_paragraph(section.heading)
        _set_heading_style(h, level=2)
        doc.add_paragraph(section.content)
        doc.add_paragraph()

    # ── STAM Scorecard Table ──────────────────────────────────────────────────
    if projects:
        h = doc.add_paragraph("STAM Scorecard Summary")
        _set_heading_style(h, level=2)

        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        for i, text in enumerate(["Project ID", "Project Name", "Type",
                                   "Score", "Classification", "Budget (ZAR)"]):
            hdr[i].text = text
            hdr[i].paragraphs[0].runs[0].bold = True

        for p in projects:
            row = table.add_row().cells
            row[0].text = str(getattr(p, "project_id", ""))
            row[1].text = str(getattr(p, "name", ""))
            row[2].text = str(getattr(p, "project_type", ""))
            row[3].text = str(getattr(p, "total_score", 0))
            row[4].text = str(getattr(p, "classification", ""))
            budget = getattr(p, "budget_rands", 0) or 0
            row[5].text = f"R{budget:,.0f}"

        doc.add_paragraph()

    # ── Risk notes ────────────────────────────────────────────────────────────
    if report.risk_notes:
        h = doc.add_paragraph("Risk and Readiness Notes")
        _set_heading_style(h, level=2)
        doc.add_paragraph(report.risk_notes)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph("─" * 70)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "Produced by STAM — Spatial Transformation Appraisal Mechanism  |  "
        "TERRA VITAL / Vastpoint  |  Gauteng Province"
    )
    footer_run.font.color.rgb = _GREY
    footer_run.font.size = Pt(8)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
